"""
Document ingestion.

Responsible for loading raw files from disk and extracting clean text,
broken down by page (PDF) or section (DOCX/TXT), so that downstream
chunking can attach accurate page/section metadata for citations.

Supported formats: PDF, DOCX, TXT.
"""
from dataclasses import dataclass, field
from pathlib import Path
import re

import pdfplumber
from docx import Document as DocxDocument

from . import config


@dataclass
class PageContent:
    """One unit of extracted text with its location within the source file."""
    text: str
    page_number: int  # 1-indexed. For DOCX/TXT this is a synthetic "section" number.
    source_file: str


@dataclass
class IngestedDocument:
    source_file: str
    pages: list = field(default_factory=list)  # list[PageContent]

    @property
    def full_text(self) -> str:
        return "\n\n".join(p.text for p in self.pages)


def _clean_text(text: str) -> str:
    """Strip common PDF artifacts: repeated whitespace, stray page-number-only
    lines, and excessive blank lines left behind by header/footer extraction."""
    if not text:
        return ""
    lines = text.split("\n")
    cleaned_lines = []
    for line in lines:
        stripped = line.strip()
        # Drop lines that are JUST a page number (typical footer artifact)
        if re.fullmatch(r"\d{1,4}", stripped):
            continue
        # Drop lines that are just a lone "Page N" / "Page N of M" footer
        if re.fullmatch(r"(?i)page\s+\d+(\s+of\s+\d+)?", stripped):
            continue
        cleaned_lines.append(line)
    text = "\n".join(cleaned_lines)
    # Collapse 3+ blank lines down to 2, collapse repeated spaces
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def ingest_pdf(path: Path) -> IngestedDocument:
    pages = []
    with pdfplumber.open(str(path)) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            raw = page.extract_text() or ""
            cleaned = _clean_text(raw)
            if cleaned:
                pages.append(PageContent(text=cleaned, page_number=i, source_file=path.name))
    return IngestedDocument(source_file=path.name, pages=pages)


def ingest_docx(path: Path) -> IngestedDocument:
    """DOCX has no fixed 'pages' (pagination is a rendering concern, not a
    structural one), so we treat each top-level Heading 1 section as a
    citation unit instead -- this is more meaningful for citation purposes
    than an arbitrary page break would be anyway."""
    doc = DocxDocument(str(path))
    pages = []
    section_num = 1
    current_section_text = []

    def flush():
        nonlocal section_num, current_section_text
        text = _clean_text("\n".join(current_section_text))
        if text:
            pages.append(PageContent(text=text, page_number=section_num, source_file=path.name))
            section_num += 1
        current_section_text = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        is_heading = para.style.name.startswith("Heading") or para.style.name == "Title"
        if is_heading and current_section_text:
            flush()
        current_section_text.append(text)
    flush()

    if not pages:  # fallback if no paragraphs detected as expected
        all_text = "\n".join(p.text for p in doc.paragraphs)
        pages = [PageContent(text=_clean_text(all_text), page_number=1, source_file=path.name)]

    return IngestedDocument(source_file=path.name, pages=pages)


def ingest_txt(path: Path) -> IngestedDocument:
    """TXT has no native section markers. We split on blank-line-delimited
    blocks that start with a Markdown-style heading-like line (all-caps or
    short title line) to approximate sections; otherwise treat the whole
    file as a single section."""
    raw = path.read_text(encoding="utf-8", errors="replace")
    cleaned = _clean_text(raw)

    # Try to split into pseudo-sections on lines that look like our own
    # section headers (see convert_docs.py: header line then '---' rule)
    blocks = re.split(r"\n-{3,}\n", cleaned)
    pages = []
    if len(blocks) > 1:
        for i, block in enumerate(blocks, start=1):
            block = block.strip()
            if block:
                pages.append(PageContent(text=block, page_number=i, source_file=path.name))
    else:
        pages = [PageContent(text=cleaned, page_number=1, source_file=path.name)]
    return IngestedDocument(source_file=path.name, pages=pages)


def ingest_file(path: Path) -> IngestedDocument:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return ingest_pdf(path)
    elif ext == ".docx":
        return ingest_docx(path)
    elif ext == ".txt":
        return ingest_txt(path)
    else:
        raise ValueError(f"Unsupported file extension: {ext}")


def ingest_directory(directory: Path) -> list:
    """Loads every supported file in `directory` (non-recursive)."""
    documents = []
    for path in sorted(directory.iterdir()):
        if path.is_file() and path.suffix.lower() in config.SUPPORTED_EXTENSIONS:
            print(f"  Ingesting {path.name} ...")
            doc = ingest_file(path)
            print(f"    -> {len(doc.pages)} page(s)/section(s), {len(doc.full_text)} characters")
            documents.append(doc)
    return documents
