"""
Converts authored markdown source documents into the final document formats
required by the assignment: PDF, DOCX, and TXT.

This script is a one-time data-prep utility, not part of the RAG application itself.
"""
import re
from pathlib import Path

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = Path(__file__).parent
OUT = SRC.parent  # data/

def parse_markdown(md_text):
    """Very small parser: returns list of (type, text) tuples."""
    blocks = []
    for line in md_text.split("\n"):
        line = line.rstrip()
        if not line.strip():
            continue
        if line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
        elif line.startswith("# "):
            blocks.append(("h1", line[2:].strip()))
        else:
            blocks.append(("p", line.strip()))
    return blocks


def to_pdf(md_path: Path, pdf_path: Path):
    text = md_path.read_text()
    blocks = parse_markdown(text)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleStyle", parent=styles["Title"], fontSize=18, spaceAfter=18, alignment=TA_LEFT)
    h2_style = ParagraphStyle("H2Style", parent=styles["Heading2"], fontSize=13, spaceBefore=14, spaceAfter=8)
    body_style = ParagraphStyle("BodyStyle", parent=styles["Normal"], fontSize=10.5, leading=15, spaceAfter=10, alignment=TA_LEFT)

    doc = SimpleDocTemplate(
        str(pdf_path), pagesize=LETTER,
        topMargin=0.9 * inch, bottomMargin=0.9 * inch,
        leftMargin=0.9 * inch, rightMargin=0.9 * inch,
        title=md_path.stem,
    )

    story = []
    for kind, content in blocks:
        content = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", content)  # bold
        if kind == "h1":
            story.append(Paragraph(content, title_style))
        elif kind == "h2":
            story.append(Paragraph(content, h2_style))
        else:
            story.append(Paragraph(content, body_style))

    def add_page_number(canvas, doc_):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.drawCentredString(LETTER[0] / 2, 0.5 * inch, f"Page {doc_.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)
    print(f"  PDF -> {pdf_path.name}")


def to_docx(md_path: Path, docx_path: Path):
    text = md_path.read_text()
    blocks = parse_markdown(text)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    for kind, content in blocks:
        content_clean = re.sub(r"\*\*(.+?)\*\*", r"\1", content)
        if kind == "h1":
            h = doc.add_heading(content_clean, level=0)
        elif kind == "h2":
            doc.add_heading(content_clean, level=1)
        else:
            doc.add_paragraph(content_clean)

    doc.save(str(docx_path))
    print(f"  DOCX -> {docx_path.name}")


def to_txt(md_path: Path, txt_path: Path):
    text = md_path.read_text()
    blocks = parse_markdown(text)
    lines = []
    for kind, content in blocks:
        content_clean = re.sub(r"\*\*(.+?)\*\*", r"\1", content)
        if kind == "h1":
            lines.append(content_clean.upper())
            lines.append("=" * len(content_clean))
        elif kind == "h2":
            lines.append("")
            lines.append(content_clean)
            lines.append("-" * len(content_clean))
        else:
            lines.append(content_clean)
        lines.append("")
    txt_path.write_text("\n".join(lines))
    print(f"  TXT -> {txt_path.name}")


if __name__ == "__main__":
    jobs = [
        ("doc1_renewable_energy.md", "pdf", "Renewable_Energy_Transition_Report.pdf"),
        ("doc2_printing_press_history.md", "docx", "Printing_Press_Information_Revolution.docx"),
        ("doc3_gut_microbiome.md", "pdf", "Gut_Microbiome_and_Human_Health.pdf"),
        ("doc4_venture_capital.md", "txt", "Startup_Fundraising_VC_Guide.txt"),
        ("doc5_hydrothermal_vents.md", "docx", "Hydrothermal_Vents_Deep_Ocean.docx"),
    ]

    print("Converting authored source documents to final formats...")
    for src_name, fmt, out_name in jobs:
        src = SRC / src_name
        out = OUT / out_name
        if fmt == "pdf":
            to_pdf(src, out)
        elif fmt == "docx":
            to_docx(src, out)
        elif fmt == "txt":
            to_txt(src, out)
    print("Done.")
